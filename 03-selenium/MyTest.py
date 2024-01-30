from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as ec
from selenium.webdriver.support.ui import WebDriverWait as webDriverWait
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import requests
from PIL import Image
import uuid
import io
import os,time
import urllib.request


headers = {
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36'
}

base_url='https://www.examtopics.com/exams/amazon/aws-certified-solutions-architect-professional-sap-c02/view/'
local_test_url = 'file:///Users/luoran/Desktop/test/Edit5.html'
simulate_img_url = '/Users/luoran/Desktop/test/image1.png'

driver = webdriver.Chrome()
driver.get(local_test_url)
document = Document()

# 测试word文件按顺序插入文本与图片
def my_function():
    list = driver.find_elements(By.XPATH, '//ul/li')
    # 遍历每个li -> 选项
    for ele in list:
        document.add_paragraph()
        # 获取div的innerHTML内容
        html_content = ele.get_attribute('innerHTML')
        # 将innerHTML转换为BeautifulSoup对象以便解析
        soup = BeautifulSoup(html_content, 'html.parser')
        # 遍历div内的所有节点
        for child in soup.children:
            if child.name == 'img':
                # 图片处理
                image_url = child['src']

                # 模拟下载图片
                img_obj = Image.open(simulate_img_url)
                ramdon_pic_name = '%s.jpg' % uuid.uuid1()
                pic_path = '03-selenium/sap_img/' + ramdon_pic_name
                
                # 将图片保存到本地以便于添加到Word文档
                with open(pic_path, 'wb') as f:
                    img_obj.save(f, format='JPEG')
                
                # 将本地图片插入到Word文档中
                last_paragraph = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                document.add_picture(pic_path, width=Inches(4))

            elif child.string and child.string.strip():  # 检查是否是文本节点且非空格
                # 文本处理
                last_paragraph = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
                last_paragraph.add_run(child.string.strip())
    
    driver.quit()
    document.save("testResult.docx")
       
# 测试从网页中下载图片, 去除ssl认证限制
def my_function2():
    
    image_url = 'https://img.examtopics.com/aws-certified-solutions-architect-professional-sap-c02/image2.png'
    response = requests.get(image_url, verify=False)
    image_file = io.BytesIO(response.content)
    img_obj = Image.open(image_file)

    ramdon_pic_name = '%s.jpg'%uuid.uuid1()
    pic_path = 'sap_img/' + ramdon_pic_name
    # 将图片保存到本地以便于添加到Word文档
    with open(pic_path, 'wb') as f:
        img_obj.save(f, format='JPEG')
    
    document = Document()
    last_paragraph = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
    document.add_picture(image_file)
    
    document.save('existing_document.docx')
    
# 测试: 点击reveal solution后获取答案列表
def my_function1():
    
    webDriverWait(driver, 100).until(
        # 此处若元素不出现需手动完成captcha验证 
        ec.presence_of_element_located((By.XPATH, '//span[@class="correct-answer"]'))
    )

    btns = driver.find_elements(By.XPATH, '//a[@class="btn btn-primary reveal-solution d-print-none"]')
    for btn in btns:
        btn.click()
    
    answers_lst = driver.find_elements(By.XPATH, '//span[@class="correct-answer"]')

    for answer in answers_lst:
        print(answer.text)
    
    driver.close()

if __name__ == "__main__":
    my_function()