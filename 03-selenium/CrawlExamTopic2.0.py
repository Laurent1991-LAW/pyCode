'''
爬取examenTopic的SAP全部题目和答案

2.0版本: 
1. 见1.0版本‘图片问题’
2. 新增操作word api
3. 改动为手动输入页码进行爬取

'''

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as ec
from selenium.webdriver.support.ui import WebDriverWait as webDriverWait
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import requests
from PIL import Image
import uuid
import io

base_url = 'https://www.examtopics.com/exams/amazon/aws-certified-solutions-architect-professional-sap-c02/view/'
index = 2
ENTER_POINT = '\r\n'

# 生成word文件
document = Document()

# 将纯文本element添加到word
def doAddTxt2Word(div_element):
    last_paragraph = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
    last_paragraph.add_run(div_element.text)
    document.add_paragraph()

# 将含图片与文本的element添加到word
def doAddImgTxt2Word(div_element):
    # 获取div的innerHTML内容 + 转换为BeautifulSoup对象以便解析
    html_content = div_element.get_attribute('innerHTML')
    soup = BeautifulSoup(html_content, 'html.parser')

    # 不存在img标签直接走纯文本逻辑
    contains_img = soup.find('img') is None
    if contains_img:
        doAddTxt2Word(div_element)
        return

    # 遍历li内的所有节点
    for child in soup.children:
        if child.name == 'img':
            # 图片处理
            image_url = child['src']

            # 下载图片
            response = requests.get(image_url, verify=False)
            image_file = io.BytesIO(response.content)
            img_obj = Image.open(image_file)

            ramdon_pic_name = '%s.jpg' % uuid.uuid1()
            pic_path = '03-selenium/sap_img/' + ramdon_pic_name
            # 将图片保存到本地以便于添加到Word文档
            with open(pic_path, 'wb') as f:
                img_obj.save(f, format='JPEG')
            # 将本地图片插入到Word文档中
            document.add_picture(pic_path, width=Inches(4))
            document.add_paragraph()

        elif child.string and child.string.strip():  # 检查是否是文本节点且非空格
            # 文本处理
            last_paragraph = document.paragraphs[-1] if document.paragraphs else document.add_paragraph()
            last_paragraph.add_run(child.string.strip())

# 遍历选项ul里的li并处理
def doProcessLiInUl(ul_element):
    li_elements = ul_element.find_elements(By.TAG_NAME, "li")
    for li_element in li_elements:
        doAddImgTxt2Word(li_element)


def doProcessExtractResult(question_titles, question_bodys, question_options_lst, answers_lst):
    for index, title in enumerate(question_titles):
        doAddTxt2Word(title)
        doAddImgTxt2Word(question_bodys[index])
        # 以ul为单位
        doProcessLiInUl(question_options_lst[index])
        doAddTxt2Word(answers_lst[index])

def crawlPage(url):
    print(f'开始爬取第[{index}]页考题及答案')
    print(url)
    driver = webdriver.Chrome()
    driver.get(url)

    webDriverWait(driver, 100).until(
        # 此处若元素不出现 需手动完成captcha验证
        ec.presence_of_element_located(
            (By.XPATH, '//span[@class="correct-answer"]'))
    )

    # 需点击reveal solution才能提取到答案
    btns = driver.find_elements(By.XPATH,
                                '//a[@class="btn btn-primary reveal-solution d-print-none"]')
    for btn in btns:
        btn.click()

    # 获取标题  /text()[not(parent::span)]
    question_titles = driver.find_elements(
        By.XPATH, '//div[@class="card-header text-white bg-primary"]')
    # 获取题干
    question_bodys = driver.find_elements(
        By.XPATH, '//div[@class="card exam-question-card"]/div/p[@class="card-text"]')
    # 获取选项
    question_options_lst = driver.find_elements(
        By.XPATH, '//div[@class="card exam-question-card"]/div/div/ul')
    # 获取正确答案
    answers_lst = driver.find_elements(
        By.XPATH, '//span[@class="correct-answer-box"]')

    doProcessExtractResult(question_titles, question_bodys,
                           question_options_lst, answers_lst)

# 获取当前页面url


def getUrl(index):
    if index > 1:
        return base_url + str(index) + '/'
    else:
        return base_url


def my_function():
    ipt = input("Please enter the page number: ")
    global index
    #while (index < 3):
    crawlPage(getUrl(int(ipt)))
    #    index += 1
    document.save(f"result{ipt}.docx")


if __name__ == "__main__":
    my_function()
